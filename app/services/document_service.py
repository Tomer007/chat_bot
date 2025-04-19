import os
import traceback
from datetime import datetime

import PyPDF2
import docx
from docx import Document
from langchain.schema import Document as LangChainDoc
from langchain.text_splitter import RecursiveCharacterTextSplitter
from werkzeug.utils import secure_filename

from app.config import UPLOAD_FOLDER
from app.utils.logging_config import logger, log_performance

# Try to import textract but don't fail if not available
try:
    import textract

    TEXTRACT_AVAILABLE = True
except ImportError:
    TEXTRACT_AVAILABLE = False
    print("Warning: textract not available. Some document formats may not be fully supported.")


@log_performance(logger)
def extract_text_from_file(file):
    """
    Extract text content from various file types (PDF, DOCX, TXT, CSV).
    Returns the extracted text as a string.
    """
    filename = file.filename.lower()
    file_content = ""

    try:
        # Handle PDF files
        if filename.endswith('.pdf'):
            logger.info(f"Processing PDF file: {filename}")
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                file_content += page.extract_text() + "\n"

        # Handle Word documents
        elif filename.endswith('.docx'):
            logger.info(f"Processing DOCX file: {filename}")
            doc = docx.Document(file)
            for para in doc.paragraphs:
                file_content += para.text + "\n"

        # Handle TXT files
        elif filename.endswith('.txt'):
            logger.info(f"Processing TXT file: {filename}")
            file_content = file.read().decode('utf-8')

        # Try textract for other file types if available
        else:
            logger.info(f"Using textract for file: {filename}")
            # Save temporarily to process with textract
            temp_path = os.path.join(UPLOAD_FOLDER, secure_filename(filename))
            file.save(temp_path)
            if TEXTRACT_AVAILABLE:
                try:
                    file_content = textract.process(temp_path).decode('utf-8')
                except Exception as e:
                    print(f"Error using textract: {str(e)}")
                    file_content = f"Could not extract text from {filename}. Unsupported file format."
            else:
                file_content = f"Could not extract text from {filename}. Textract not available for this file format."

            # Clean up temp file
            if os.path.exists(temp_path):
                os.remove(temp_path)

        return file_content.strip()

    except Exception as e:
        logger.error(f"Error extracting text from {filename}: {str(e)}")
        return f"Error extracting text: {str(e)}"


@log_performance(logger)
def load_docx(path: str):
    """
    Load content from .docx or .txt files and return as a LangChain Document.
    """
    logger.info(f"Loading document from: {path}")
    try:
        file_ext = os.path.splitext(path)[1].lower()

        # Check if file exists
        if not os.path.isfile(path):
            logger.error(f"Document not found: {path}")
            return LangChainDoc(page_content="This is a default reference text. The document could not be found.")

        # Get file size
        file_size = os.path.getsize(path)
        logger.debug(f"Document size: {file_size / 1024:.2f}KB")

        if file_ext == '.docx':
            doc = Document(path)
            para_count = len(doc.paragraphs)
            empty_para_count = sum(1 for para in doc.paragraphs if not para.text.strip())
            full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip() != ""])

            logger.debug(f"Loaded {para_count} paragraphs from DOCX ({empty_para_count} empty)")
            logger.info(f"Document loaded: {len(full_text)} characters")

        elif file_ext == '.txt':
            with open(path, 'r', encoding='utf-8') as f:
                full_text = f.read()

            line_count = full_text.count('\n') + 1
            word_count = len(full_text.split())
            logger.debug(f"Loaded {line_count} lines, {word_count} words from TXT file")
            logger.info(f"Document loaded: {len(full_text)} characters")

        else:
            logger.error(f"Unsupported file type: {file_ext}")
            raise ValueError(f"Unsupported file type: {file_ext}")

        return LangChainDoc(page_content=full_text)

    except Exception as e:
        error_details = traceback.format_exc()
        logger.error(f"Error loading document: {str(e)}\n{error_details}")
        # Return a default document with some basic information
        return LangChainDoc(page_content="This is a default reference text. The actual document could not be loaded.")


@log_performance(logger)
def load_and_chunk_document(file_path: str, chunk_size: int = 1500, chunk_overlap: int = 200):
    """
    Load a document and split it into chunks for processing.
    """
    logger.info(f"Loading and chunking document: {file_path} (chunk_size={chunk_size}, overlap={chunk_overlap})")
    doc = load_docx(file_path)

    # Split if needed for large files
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    doc_chunks = splitter.split_documents([doc])

    # Calculate average chunk size
    total_chars = sum(len(chunk.page_content) for chunk in doc_chunks)
    avg_chunk_size = total_chars / len(doc_chunks) if doc_chunks else 0

    logger.info(f"Document split into {len(doc_chunks)} chunks (avg size: {avg_chunk_size:.2f} chars)")

    # Log detailed info about chunks
    for i, chunk in enumerate(doc_chunks):
        logger.debug(f"Chunk {i + 1}/{len(doc_chunks)}: {len(chunk.page_content)} chars")

    # Return all chunks
    return doc_chunks


def process_uploaded_file(file):
    """Process an uploaded file, extract text, and return file information"""
    if not file or not file.filename:
        return {}, ""

    try:
        # Get basic file information
        filename = secure_filename(file.filename)
        file_extension = os.path.splitext(filename)[1].lower()
        file_size = 0

        # Create upload folder if it doesn't exist
        if not os.path.exists(UPLOAD_FOLDER):
            os.makedirs(UPLOAD_FOLDER)

        # Save the file temporarily to get its size and process it
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(file_path)
        file_size = os.path.getsize(file_path)

        # Extract text content
        file.seek(0)  # Reset file pointer
        text_content = extract_text_from_file(file)

        # Prepare file info dictionary
        file_info = {
            'filename': filename,
            'size': file_size,
            'type': file_extension,
            'uploaded_at': datetime.now().isoformat(),
            'text_length': len(text_content) if text_content else 0
        }

        # Log the file processing
        logger.info(f"Processed file: {filename}, size: {file_size} bytes, extracted text: {len(text_content)} chars")

        # Clean up temporary file if needed
        if os.path.exists(file_path):
            os.remove(file_path)

        return file_info, text_content

    except Exception as e:
        error_message = f"Error processing file {file.filename if file else 'unknown'}: {str(e)}"
        logger.error(error_message)
        return {'error': error_message}, ""
