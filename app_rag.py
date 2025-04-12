import csv
import io
import json
import os
import tempfile
from datetime import datetime, timedelta
from functools import wraps

import docx
import pdfplumber
import yaml
from docx import Document
from dotenv import load_dotenv
from flask import Flask, render_template, request, jsonify, session, redirect, url_for
from flask_cors import CORS
from langchain.schema import Document as LangChainDoc
from langchain.text_splitter import RecursiveCharacterTextSplitter
from openai import OpenAI
from werkzeug.utils import secure_filename

load_dotenv()
app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY", "supersecret")
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(minutes=5)  # Set session timeout to 3 minutes
CORS(app)


def load_chatbot_config():
    # Ensure the path is correct relative to your app file
    config_path = os.path.join(os.path.dirname(__file__), 'config', 'chatbot_config.yaml')

    with open(config_path, 'r', encoding='utf-8') as file:
        config_data = yaml.safe_load(file)

    chatbots = config_data.get("chatbots", {})

    # Get the environment variable, defaulting to "default"
    selected_chatbot = os.getenv("CHATBOT_NAME", "default").strip()

    # If not found by key, try searching by the chatbot's "name" field
    for key, config in chatbots.items():
        if config.get("name") == selected_chatbot:
            # Look up the configuration within the "chatbots" dictionary
            chosen_config = config
            break

    if not chosen_config:
        # Log a warning and fallback to default if the selected key isn't found
        chosen_config = config_data.get("chatbots", {}).get("default", {})

    return chosen_config


# Example of calling the function:
chatbot_config = load_chatbot_config()

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


# Login required decorator
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)

    return decorated_function


# --- Load external content from .docx or .txt ---
def load_docx(path: str):
    try:
        file_ext = os.path.splitext(path)[1].lower()
        if file_ext == '.docx':
            doc = Document(path)
            full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip() != ""])
        elif file_ext == '.txt':
            with open(path, 'r', encoding='utf-8') as f:
                full_text = f.read()
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

        return LangChainDoc(page_content=full_text)
    except Exception as e:
        print(f"Error loading document: {e}")
        # Return a default document with some basic information
        return LangChainDoc(page_content="This is a default reference text. The actual document could not be loaded.")


# Load and optionally chunk the document
docx_file = os.getenv("DOCX_FILE_PATH", "doc/rag_content.txt")

pdn_doc = load_docx(docx_file)

# Optional: split if needed for large files
splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=200)
doc_chunks = splitter.split_documents([pdn_doc])
# Take top 1–2 chunks if needed (can be improved with retrieval later)
reference_text = "\n\n".join([chunk.page_content for chunk in doc_chunks[:2]])

# Get the prompt file path from an environment variable
prompt_file = os.getenv("BUDDY_PROMPT_FILE", "prompts/buddy_prompt.txt")

# Read the prompt from the specified file
with open(prompt_file, "r", encoding="utf-8") as f:
    base_prompt = f.read().strip()

SYSTEM_PROMPT = f"""{base_prompt}

להלן מידע נוסף :

{reference_text}
"""

# Chat history per session
conversation_history = {}


def get_session_id():
    if 'session_id' not in session:
        session['session_id'] = os.urandom(8).hex()
        print(f"\n🆕 New session created: {session['session_id']}")
    else:
        print(f"\n📝 Using existing session: {session['session_id']}")
    return session['session_id']


def generate_response(user_message):
    session_id = get_session_id()

    if session_id not in conversation_history:
        conversation_history[session_id] = [{"role": "system", "content": SYSTEM_PROMPT}]
        print(f"\n📚 Initialized conversation history for session: {session_id}")
    else:
        print(f"\n📖 Retrieved existing conversation history for session: {session_id}")

    history = conversation_history[session_id]
    history.append({"role": "user", "content": user_message})

    print("\n📥 Prompt sent to GPT:\n----------------------" + user_message)
    print(json.dumps(history, indent=2, ensure_ascii=False))

    try:
        response = client.chat.completions.create(
            model="gpt-4",
            messages=history,
            temperature=0.7
        )
        bot_reply = response.choices[0].message.content.strip()
        history.append({"role": "assistant", "content": bot_reply})
    except Exception as e:
        bot_reply = f"I'm sorry, I encountered an error: {e}"

    return bot_reply


@app.before_request
def before_request():
    if 'user_id' in session:
        # Check if session is expired
        if 'last_activity' in session:
            last_activity = datetime.fromisoformat(session['last_activity'])
            if datetime.now() - last_activity > app.config['PERMANENT_SESSION_LIFETIME']:
                session.clear()
                return redirect(url_for('login'))
        # Update last activity time
        session['last_activity'] = datetime.now().isoformat()


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form.get("username")
        password = request.form.get("password")
        remember = request.form.get("remember") == "on"
        print("password " + password)
        if password == os.getenv("USER_PASSWORD"):
            session.permanent = remember
            session["user_id"] = username
            next_page = request.args.get("next")
            if next_page and next_page.startswith("/"):
                return redirect(next_page)
            return redirect(url_for("index"))
        else:
            return render_template("login.html", error="שם משתמש או סיסמה שגויים", chatbot_config=load_chatbot_config())

    # GET request - show login form
    return render_template("login.html", chatbot_config=load_chatbot_config())


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for('login'))


@app.route("/")
@login_required
def index():
    return render_template("chat.html", chatbot_config=chatbot_config)


@app.route("/session-status", methods=["GET"])
@login_required
def session_status():
    session_id = get_session_id()
    return jsonify({
        "session_id": session_id,
        "has_history": session_id in conversation_history,
        "history_length": len(conversation_history.get(session_id, []))
    })


@app.route('/check_session')
def check_session():
    if 'user_id' in session and 'last_activity' in session:
        last_activity = datetime.fromisoformat(session['last_activity'])
        if datetime.now() - last_activity > app.config['PERMANENT_SESSION_LIFETIME']:
            session.clear()
            return jsonify({'status': 'expired'})
        return jsonify({'status': 'active'})
    return jsonify({'status': 'expired'})


def extract_text_from_file(file):
    """
    Extract text content from various file types (PDF, DOCX, TXT, CSV).
    Returns the extracted text as a string.
    """
    filename = secure_filename(file.filename)
    file_ext = os.path.splitext(filename)[1].lower()

    # ----- PDF Handling -----
    if file_ext == '.pdf':
        temp_file_path = None
        try:
            # 1) Save the uploaded PDF to a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                file.save(tmp.name)
                temp_file_path = tmp.name

            # 2) Open and read the PDF using pdfplumber
            extracted_text = ""
            with pdfplumber.open(temp_file_path) as doc:
                for page in doc.pages:
                    page_text = page.extract_text()
                    if page_text:
                        extracted_text += page_text + "\n"
            return extracted_text

        finally:
            # 3) Clean up the temp file
            if temp_file_path and os.path.exists(temp_file_path):
                os.remove(temp_file_path)

    # ----- DOCX Handling -----
    elif file_ext == '.docx':
        doc = docx.Document(io.BytesIO(file.read()))
        paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        return "\n".join(paragraphs)

    # ----- TXT Handling -----
    elif file_ext == '.txt':
        # Assume UTF-8 decoding. Adjust as needed (e.g. 'latin-1' or 'charmap')
        content = file.read().decode('utf-8', errors='replace')
        return content

    # ----- CSV Handling -----
    elif file_ext == '.csv':
        content = file.read().decode('utf-8', errors='replace')
        csv_reader = csv.reader(io.StringIO(content))
        # Join each row with commas and each row with a newline
        return "\n".join([",".join(row) for row in csv_reader])

    else:
        raise ValueError(f"Unsupported file extension: {file_ext}")


@app.route('/api/chat_upload', methods=['POST'])
def handle_input():
    """Accept both user text and optional file, combine, and send to GPT."""
    user_message = request.form.get("message", "").strip()
    uploaded_file = request.files.get("file")  # None if not provided
    file_content = ""
    if uploaded_file and uploaded_file.filename:
        try:
            uploaded_file.seek(0)  # Reset pointer
            file_content = extract_text_from_file(uploaded_file)
            print(f"📝 Extracted text length: {len(file_content)} characters")
            print("file_content = " + file_content)
        except Exception as e:
            print(f"❌ Error processing file: {e}")
            return jsonify({"success": False, "error": str(e)}), 500

    # Combine text + file content into a single prompt
    # Adjust the text below to your needs
    combined_prompt = (
        f"{user_message if user_message != '' else 'Please analyze and create a PDN report based on the content above.'}\n{file_content}\n"
    )

    try:
        # Send combined prompt to GPT
        response_text = generate_response(combined_prompt)
        return jsonify({"success": True, "response": response_text})
        # return ""

    except Exception as e:
        print(f"❌ Error generating response: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


def read_docx(file):
    """Read content from a docx file."""
    doc = Document(file)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])


if __name__ == "__main__":
    app.run(host='0.0.0.0', port=5001, debug=True)
